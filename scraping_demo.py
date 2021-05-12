# Web Scraping Demo

import time
import os
import string
from datetime import datetime
import requests
from diskcache import Cache
from bs4 import BeautifulSoup
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor

class Fox():
    """
    A wrapper for requests that automates interaction with diskcache
    """

    def __init__(self):
        self.base = None
        self.params = None
        self.headers = []
        self.private_keys = []
        self.cache_path = None
        self.accept_codes = [200]

    def update(self,
        base = None,
        params = None,
        headers = None,
        private_keys = None,
        cache_path = None,
        cache_ref = None,
        accept_codes = None,
    ):

        """
        Used to set up conditions for a request, including base, parameters, cache reference, and which codes to accept
        """

        if base != None:
            self.base = base
            self.params = None
        if params != None:
            self.params = params
        if headers != None:
            self.headers = headers
        if private_keys != None:
            self.private_keys = private_keys
        if cache_path != None:
            self.cache_path = cache_path
            if not os.path.exists(self.cache_path):
                os.makedirs(self.cache_path)
        if cache_ref != None:
            self.cache_ref = cache_ref
        if accept_codes != None:
            self.accept_codes = accept_codes

        self.cache_key = self.make_cache_key(self.base,self.params,self.private_keys)

    def make_cache_key(self, base=None, params=None, private_keys=None):
        """
        Makes a unique string for cache access
        """

        kvs = []
        if (base == None) or (params in [None,{}]):
            return base
        else:
            alpha_keys = sorted(params.keys())
            for k in alpha_keys:
                if k not in private_keys:
                    kvs.append('{}-{}'.format(k, params[k]))
            return base + '_'.join(kvs)

    def make_request(self, sleep = None):

        """
        Handles whether to actually send the request or just skip it because it's already cached
        """

        if self.base == None:
            return None


        returnable = None

        if self.cache_ref == None:

            if self.cache_path == None:
                response = send_request()

            else:
                with Cache(self.cache_path) as cache:
                    if self.cache_key not in cache:
                        returnable = self.send_request(cache)

        else:

            if self.cache_key not in self.cache_ref:
                if sleep != None:
                    time.sleep(sleep)
                returnable = self.send_request(self.cache_ref)

        if returnable != None:
            return returnable

    def send_request(self, cache=None):

        """
        Actually sends the request and checks the response code
        """

        response = requests.get(
            self.base,
            params=self.params,
            headers=self.headers,
        )

        if response.status_code not in self.accept_codes:
            print("Response:",response.status_code,':\n',self.base,self.params,self.headers)
            print(response.text)
        else:
            if cache != None:
                cache[self.cache_key] = response

        return response

def manage_scraping(row,cache,fox):

    description_found = False

    for isbn_format in ["ebook","hardcover","paper"]:

        if "University of California Press" == row["publisher"]:

            if description_found == False:

                isbn = row[isbn_format]
                url = "https://www.ucpress.edu/book/" + str(isbn)
                fox.update(base=url)
                fox.make_request(sleep=3)
                resp = cache[fox.cache_key]
                if resp.status_code != 200:
                    print("Nothing found at", url)
                    print(resp.status_code)
                else:
                    print("Webpage found at", url)
                    soup = BeautifulSoup(resp.text,"html.parser")

                    description_section = soup.find("section",id="link-about-book")

                    if description_section != None:
                        description_found = True
                        just_the_description = description_section.find("article")
                        simplified_text = text_with_newlines(just_the_description)
                        row["description"] = simplified_text

    return row

def manage_apply():
    spreadsheet = pd.read_excel("metadata.xlsx",dtype=str)

    new_headers = [
        "handle",
        "title",
        "subtitle",
        "publisher",
        "ebook",
        "hardcover",
        "paper",
        "description"
    ]

    spreadsheet = spreadsheet.reindex(columns = new_headers)

    with Cache("demo_cache") as cache:
        user_agent = ""
        fox = Fox()
        fox.update(
            headers = {'User-Agent':user_agent},
            cache_ref = cache,
            accept_codes=[200]
        )

        spreadsheet = spreadsheet.apply(
                lambda row : manage_scraping(row,cache,fox),
                axis=1
            )

    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")

    outdir = "outputs"
    if not os.path.exists(outdir):
        os.makedirs(outdir)
    spreadsheet.to_excel(f"{outdir}/{timestamp}_output.xlsx",index=False)

    new_document = Document()
    new_document.add_heading("Scraped Book Descriptions",0)
    new_document.styles["Normal"].font.name = "Times New Roman"
    new_document.styles["Heading 1"].font.size = Pt(10)
    for each_style in ["Title","Heading 1","Heading 2"]:
        new_document.styles[each_style].font.name = None
        new_document.styles[each_style].font.color.rgb = RGBColor(0,0,0)

    for i in spreadsheet.index.values:
        if not pd.isnull(spreadsheet.loc[i,"description"]):
            new_document.add_heading(spreadsheet.loc[i,"handle"],level=1)
            full_title = spreadsheet.loc[i,"title"] + ": " + spreadsheet.loc[i,"subtitle"]
            new_document.add_heading(full_title,level=2)
            new_document.add_paragraph(spreadsheet.loc[i,"description"])

    new_document.save(f"{outdir}/{timestamp}_output.docx")
    print("Output files created")

def text_with_newlines(elem):

    '''
    A custom alternative to BeautifulSoup's string methods that keeps line breaks represented by div and br elements. Whitespace and line breaks in JADE transcriptions are often used to represent spatial distance in the analog page.
    Credit: https://gist.github.com/zmwangx/ad0830ba94b1fd98f428
    '''

    still_inside_em = False
    text = ''
    for e in elem.descendants:
        if isinstance(e, str):
            if not still_inside_em:
                text += str(e)
            still_inside_em = False
        elif e.name in ['em','i']:
            text += str(e)
            still_inside_em = True
        elif e.name in ['br', 'p', 'div','article']:
            text += '\n'

    return text.replace("\n\n","\n").strip("\n"+string.whitespace)

if __name__ == '__main__':
    manage_apply()
